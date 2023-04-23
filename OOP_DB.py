import re
import string
import nltk
import pymorphy2
import json
import folium
from folium import plugins
from folium import FeatureGroup
from folium.plugins import MarkerCluster
from nltk.corpus import stopwords
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Mm
from fuzzywuzzy import fuzz
import os


class TeleGOD(object):

    def __init__(self, DbAdvanced):

        self.total_posts = None

        self.spec_chars = string.punctuation + r'\n\x0«»\t—…[]\n*'
        self.stop_words = stopwords.words('russian')
        self.morph = pymorphy2.MorphAnalyzer()

        ### Используя JSON файл получаю названия 407 городов с их координатами
        with open("ua-cities.json", encoding="utf8") as f:
            d = json.load(f)

        self.city_list = []
        self.city_lon = []
        self.city_lat = []

        for region in d[0]['regions']:
            for item in region['cities']:
                self.city_list.append(item['name'].lower())
                self.city_lon.append(item['lat'].lower())
                self.city_lat.append(item['lng'].lower())

        city_info = list(zip(self.city_list, self.city_lon, self.city_lat))

        self.cities_dict = dict()
        self.strikes_dict = dict()
        self.reports_dict = dict()
        self.tag_dict = dict()

        for index in range(len(city_info)):
            self.cities_dict[str(self.city_list[index])] = []
            self.strikes_dict[str(self.city_list[index])] = []
            self.reports_dict[str(self.city_list[index])] = []
            self.tag_dict[str(self.city_list[index])] = []

        self.total_strikes = None
        self.filtered_reports_dict = None
        self.cities_cleaned_dict, self.strikes_cleaned_dict = None, None
        self.last_day = None
        self.is_build = False

        self.DB = DbAdvanced

    ### В данном методе обрабатывается каждое сообщение (удаление спец. символов и стоп слов, токенизация и лематизация текста).
    ### Затем создаются словари для дальнейшего отображения на карте.
    def main_actions(self, dataset):

        total_posts = 0
        total_strikes = 0

        for index in range(len(dataset)):

            adv_text = json.loads(dataset[index]['ADV_MESSAGE'])

            for key in self.cities_dict:

                if key in adv_text and (all(['хроника' not in adv_text, 'сводка' not in adv_text,
                                               'карта' not in adv_text, 'обстановка' not in adv_text,
                                               'направление' not in adv_text])):

                    link = '<a href=' + str(dataset[index]['SENDER']) + '/' + str(dataset[index]['MESSAGE_ID']) + '>'\
                           + str(dataset[index]['SENDER']) + '/' + str(dataset[index]['MESSAGE_ID']) + '</a>'

                    message_and_link = [dataset[index]['MESSAGE'], link]

                    self.cities_dict[str(key)].append(message_and_link)

                    total_posts += 1

                    self.reports_dict[str(key)].append(str(dataset[index]['MESSAGE']))

                if key in adv_text and (all(['хроника' not in adv_text, 'сводка' not in adv_text, 'карта' not in adv_text,
                     'обстановка' not in adv_text, 'направление' not in adv_text])) and ('удар' in adv_text or
                                                                                             'обстрел' in adv_text or
                                                                                             'обстрелять' in adv_text or
                                                                                             'прилёт' in adv_text or
                                                                                             'прилет' in adv_text or
                                                                                             'взрыв' in adv_text):
                    link = '<a href=' + str(dataset[index]['SENDER']) + '/' + str(dataset[index]['MESSAGE_ID']) + '>' \
                           + str(dataset[index]['SENDER']) + '/' + str(dataset[index]['MESSAGE_ID']) + '</a>'

                    message_and_link = [dataset[index]['MESSAGE'], link]

                    self.strikes_dict[str(key)].append(message_and_link)

                    total_strikes += 1

        return self.cities_dict, self.strikes_dict, total_posts, total_strikes, self.reports_dict

    def dict_cleaning(self, reports_dict, uniq):

        cleaned_dict = dict()

        for key in reports_dict:
            cleaned_dict[key] = []
            for i in range(len(reports_dict[key])-1):
                for j in range(i + 1, len(reports_dict[key])):
                    if fuzz.token_set_ratio(reports_dict[key][i], reports_dict[key][j]) > 100 - uniq:
                        reports_dict[key][j] = '_'

        for key in reports_dict:
            for i in range(len(reports_dict[key])):
                if reports_dict[key][i] != '_':
                    cleaned_dict[key].append(reports_dict[key][i])

        return cleaned_dict

    def cities_and_strikes_cleaning(self, cities_dict, strikes_dict, uniq):

        cities_cleaned_dict = dict()

        for key in cities_dict:
            cities_cleaned_dict[key] = []
            for i in range(len(cities_dict[key])-1):
                for j in range(i + 1, len(cities_dict[key])):
                    if fuzz.token_set_ratio(cities_dict[key][i][0], cities_dict[key][j][0]) > 100 - uniq:
                        cities_dict[key][j][0] = '_'

        for key in cities_dict:
            for i in range(len(cities_dict[key])):
                if cities_dict[key][i][0] != '_':
                    cities_cleaned_dict[key].append(cities_dict[key][i][1])

        strikes_cleaned_dict = dict()

        for key in strikes_dict:
            strikes_cleaned_dict[key] = []
            for i in range(len(strikes_dict[key])-1):
                for j in range(i + 1, len(strikes_dict[key])):
                    if fuzz.token_set_ratio(strikes_dict[key][i][0], strikes_dict[key][j][0]) > 100 - uniq:
                        strikes_dict[key][j][0] = '_'

        for key in strikes_dict:
            for i in range(len(strikes_dict[key])):
                if strikes_dict[key][i][0] != '_':
                    strikes_cleaned_dict[key].append(strikes_dict[key][i][1])

        return cities_cleaned_dict, strikes_cleaned_dict

    def report(self, filtered_reports_dict, begin, end, uniq):

        begin_report = datetime.utcfromtimestamp(begin + 86400).strftime('%Y-%m-%d')
        end_report = datetime.utcfromtimestamp(end + 86400).strftime('%Y-%m-%d')
        document = Document()
        style = document.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        section = document.sections[0]
        section.right_margin = Mm(10)
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(10)
        section.header_distance = Mm(10)
        section.footer_distance = Mm(10)
        head = document.add_heading('Отчет по оперативной обстановке\n {} - {}\n Уникальность информации - {}%'.format(str(begin_report), str(end_report), str(uniq)))
        head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for key in filtered_reports_dict:
            if len(filtered_reports_dict[key]) != 0:

                par = document.add_paragraph().add_run(key.capitalize())
                par.font.size = Pt(14)
                par.bold = True

                for point in range(len(filtered_reports_dict[key])):
                    par0 = document.add_paragraph(str(point+1) + ') ')
                    par0.add_run(filtered_reports_dict[key][point].strip())
                    par0.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    fmt0 = par0.paragraph_format
                    fmt0.first_line_indent = Mm(15)

        #report_name = r'C:\Users\Chubu\OneDrive\Documents\GitHub\GeoNews\reports\report_' + str(begin_report) + '-' + str(end_report) + '_' + str(uniq) + '.docx'
        report_name = 'report.docx'
        document.save(report_name)
        os.startfile(report_name)

    def map_creation(self, filtered_cities_dict, filtered_strikes_dict, total_posts, total_strikes, begin, end, uniq): ### Создание карты и отображение на ней меток

        map = folium.Map(width=1300,
                         height=780,
                         location=[49.05, 30.60],
                         tiles='openstreetmap',
                         zoom_start=6,
                         min_zoom=1,
                         max_zoom=14)

        plugins.Geocoder().add_to(map)

        fmtr = "function(num) {return L.Util.formatNum(num, 3) + ' º ';};"
        plugins.MousePosition(
            position="topright",
            separator=" | ",
            prefix="Coordinates:",
            lat_formatter=fmtr,
            lng_formatter=fmtr).add_to(map)

        minimap = plugins.MiniMap()
        map.add_child(minimap)

        plugins.Fullscreen().add_to(map)

        # plugins.LocateControl().add_to(map)

        plugins.MeasureControl(position='topright',
                               primary_length_unit='meters',
                               secondary_length_unit='miles',
                               primary_area_unit='sqmeters',
                               secondary_area_unit='acres').add_to(map)

        folium.TileLayer('Stamen Toner').add_to(map)
        folium.TileLayer('Stamen Terrain').add_to(map)
        folium.TileLayer('Stamen Watercolor').add_to(map)
        folium.TileLayer('openstreetmap').add_to(map)
        folium.TileLayer('cartodbpositron').add_to(map)
        folium.TileLayer('cartodbdark_matter').add_to(map)

        folium.TileLayer(
            tiles='https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}',
            attr='Esri',
            name='Esri Satellite',
            overlay=False,
            control=True
        ).add_to(map)

        plugins.Draw(
            # export=True,
            # filename="my_data.geojson",
            # position="topleft",
            # draw_options={"polyline": {"allowIntersection": False}},
            # edit_options={"poly": {"allowIntersection": False}},
        ).add_to(map)

        marker_cluster = MarkerCluster(name='Clusters').add_to(map)
        marker_points = FeatureGroup(name='All markers', show=False).add_to(map)
        strikes_point = FeatureGroup(name='Strikes', show=False).add_to(map)

        for i in range(0, len(filtered_cities_dict)):

            if len(filtered_cities_dict[self.city_list[i]]) != 0:

                folium.Marker(location=[self.city_lon[i], self.city_lat[i]],
                              popup=filtered_cities_dict[self.city_list[i]],
                              tooltip=[self.city_list[i].capitalize(), len(filtered_cities_dict[self.city_list[i]]), 'новость(и)'],
                              icon=folium.Icon(color='orange' if (len(filtered_cities_dict[self.city_list[i]])) > 4 else 'green',
                                               icon="info-sign")).add_to(marker_cluster)

                folium.Marker(location=[self.city_lon[i], self.city_lat[i]],
                              popup=filtered_cities_dict[self.city_list[i]],
                              tooltip=[self.city_list[i].capitalize(), len(filtered_cities_dict[self.city_list[i]]), 'новость(и)'],
                              icon=folium.Icon(color='orange' if (len(filtered_cities_dict[self.city_list[i]])) > 4 else 'green',
                                               icon="info-sign")).add_to(marker_points)

                folium.Circle(radius=((len(filtered_cities_dict[self.city_list[i]])) / total_posts) * 50000,
                              location=(self.city_lon[i], self.city_lat[i]),
                              color='orange',
                              weiht=5,
                              fill=True,
                              fill_color='orange' if (len(filtered_cities_dict[self.city_list[i]])) > 4 else 'orange',
                              fill_opacity=0.5,
                              stroke=False
                              ).add_to(marker_points)

            if len(filtered_strikes_dict[self.city_list[i]]) != 0:

                folium.Marker(location=[float(self.city_lon[i]) - 0.02, float(self.city_lat[i]) + 0.02],
                              popup=filtered_strikes_dict[self.city_list[i]],
                              tooltip=[self.city_list[i].capitalize(), len(filtered_strikes_dict[self.city_list[i]]), 'удар(а)'],
                              icon=folium.Icon(color='red', icon="fire")).add_to(marker_points)

                folium.Circle(radius=((len(filtered_cities_dict[self.city_list[i]])) / total_strikes) * 50000,
                              location=(float(self.city_lon[i]) - 0.02, float(self.city_lat[i]) + 0.02),
                              color='orange',
                              weiht=5,
                              fill=True,
                              fill_color='red',
                              fill_opacity=0.5,
                              stroke=False
                              ).add_to(marker_points)

                folium.Marker(location=[float(self.city_lon[i]) - 0.02, float(self.city_lat[i]) + 0.02],
                              popup=filtered_strikes_dict[self.city_list[i]],
                              tooltip=[self.city_list[i].capitalize(), len(filtered_strikes_dict[self.city_list[i]]), 'удар(а)'],
                              icon=folium.Icon(color='red', icon="fire")).add_to(strikes_point)

                folium.Circle(radius=((len(filtered_cities_dict[self.city_list[i]])) / total_strikes) * 50000,
                              location=(float(self.city_lon[i]) - 0.02, float(self.city_lat[i]) + 0.02),
                              color='orange',
                              weiht=5,
                              fill=True,
                              fill_color='red',
                              fill_opacity=0.5,
                              stroke=False
                              ).add_to(strikes_point)

        folium.LayerControl().add_to(map)

        begin_map = datetime.utcfromtimestamp(begin+86400).strftime('%Y-%m-%d')
        end_map = datetime.utcfromtimestamp(end+86400).strftime('%Y-%m-%d')

        #map_name = 'map_' + str(begin_map) + '-' + str(end_map) + '_' + str(uniq) + '.html'
        map_name = 'map.html'
        map.save(map_name)

        with open(map_name, 'r', encoding="utf8") as f:
            self.html = f.read()

    def tag_map_creation(self, tag_word, begin, end):

        self.tag_word = self.morph.parse(tag_word)[0].normal_form

        self.total_tag_words = 0

        self.tag_dict = dict()
        city_info = list(zip(self.city_list, self.city_lon, self.city_lat))

        for index in range(len(city_info)):
            self.tag_dict[str(self.city_list[index])] = []

        self.last_day = self.DB.read_db(begin, end)

        # self.features = []

        for index in range(len(self.last_day)):

            text = json.loads(self.last_day[index]['ADV_MESSAGE'])
            # date = datetime.utcfromtimestamp(self.last_day[index]['DATE']).strftime('%Y-%m-%d %H:%M:%S')

            for key in self.tag_dict:
                if key in text and self.tag_word in text and (all(['хроника' not in text,
                                                                  'сводка' not in text,
                                                                  'карта' not in text,
                                                                  'обстановка' not in text,
                                                                  'направление' not in text])):

                    link = '<a href=' + str(self.last_day[index]['SENDER']) + '/' \
                           + str(self.last_day[index]['MESSAGE_ID']) + '>' + str(self.last_day[index]['SENDER'])\
                           + '/' + str(self.last_day[index]['MESSAGE_ID']) + '</a>'

                    self.tag_dict[str(key)].append(link)

                    self.total_tag_words += 1

                    # feature = {
                    #     'type': 'Feature',
                    #     'geometry': {
                    #         'type': 'Point',
                    #         'coordinates': [city_info[index][1], city_info[index][2]]
                    #     },
                    #     'properties': {
                    #         'time': self.last_day[index]['DATE'],
                    #         "tooltip": "my tooltip text"
                    #         'style': {'color': ''},
                    #         'icon': 'circle',
                    #         'iconstyle': {
                    #             'fillColor': '#0000FF',
                    #             'fillOpacity': 0.8,
                    #             'stroke': 'true',
                    #             'radius': 5
                    #
                    #         }
                    #     }
                    # }
                    # self.features.append(feature)

        tag_map = folium.Map(width=1300,
                             height=780,
                             location=[49.05, 30.60],
                             tiles='openstreetmap',
                             zoom_start=6,
                             min_zoom=1,
                             max_zoom=14)

        # plugins.TimestampedGeoJson(self.features,
        #                            period='P1D',
        #                            duration='P1H',
        #                            transition_time=1000,
        #                            auto_play=True).add_to(tag_map)

        # plugins.TimestampedGeoJson(
        #     {"type": "FeatureCollection", "features": features},
        #     period="P1D",
        #     add_last_point=True,
        #     auto_play=False,
        #     loop=False,
        #     max_speed=1,
        #     loop_button=True,
        #     #date_options="YYYY/MM/DD",
        #     time_slider_drag_update=True,
        #     duration="P1H",
        # ).add_to(tag_map)

        plugins.Geocoder().add_to(tag_map)

        fmtr = "function(num) {return L.Util.formatNum(num, 3) + ' º ';};"
        plugins.MousePosition(
            position="topright",
            separator=" | ",
            prefix="Coordinates:",
            lat_formatter=fmtr,
            lng_formatter=fmtr).add_to(tag_map)

        minimap = plugins.MiniMap()
        tag_map.add_child(minimap)

        plugins.Fullscreen().add_to(tag_map)

        # plugins.LocateControl().add_to(map)

        plugins.MeasureControl(position='topright',
                               primary_length_unit='meters',
                               secondary_length_unit='miles',
                               primary_area_unit='sqmeters',
                               secondary_area_unit='acres').add_to(tag_map)

        folium.TileLayer('Stamen Toner').add_to(tag_map)
        folium.TileLayer('Stamen Terrain').add_to(tag_map)
        folium.TileLayer('Stamen Watercolor').add_to(tag_map)
        folium.TileLayer('openstreetmap').add_to(tag_map)
        folium.TileLayer('cartodbpositron').add_to(tag_map)
        folium.TileLayer('cartodbdark_matter').add_to(tag_map)

        folium.TileLayer(
            tiles='https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}',
            attr='Esri',
            name='Esri Satellite',
            overlay=False,
            control=True
        ).add_to(tag_map)

        plugins.Draw(
            # export=True,
            # filename="my_data.geojson",
            # position="topleft",
            # draw_options={"polyline": {"allowIntersection": False}},
            # edit_options={"poly": {"allowIntersection": False}},
        ).add_to(tag_map)

        folium.LayerControl().add_to(tag_map)

        for i in range(len(self.tag_dict)):

            if len(self.tag_dict[self.city_list[i]]) != 0:

                folium.Marker(location=[self.city_lon[i], self.city_lat[i]],
                              popup=self.tag_dict[self.city_list[i]],
                              tooltip=[self.city_list[i].capitalize(), len(self.tag_dict[self.city_list[i]]), str(tag_word)],
                              icon=folium.Icon(color='blue', icon="info-sign")).add_to(tag_map)

                # print(self.tag_dict[self.city_list[i]])

                folium.Circle(radius=((len(self.tag_dict[self.city_list[i]])) / self.total_tag_words) * 50000,
                              location=(float(self.city_lon[i]), float(self.city_lat[i])),
                              color='blue',
                              weiht=5,
                              fill=True,
                              fill_color='blue',
                              fill_opacity=0.5,
                              stroke=False
                              ).add_to(tag_map)

        tag_map_name = 'tag_map.html'
        tag_map.save(tag_map_name)

        with open(tag_map_name, 'r', encoding="utf8") as ff:
            self.tag_html = ff.read()

    def calculate(self, begin, end, uniq):

        self.total_strikes = None
        self.filtered_reports_dict = None
        self.cities_cleaned_dict, self.strikes_cleaned_dict = None, None
        self.last_day = None
        self.is_build = False

        self.cities_dict = dict()
        self.strikes_dict = dict()
        self.reports_dict = dict()

        city_info = list(zip(self.city_list, self.city_lon, self.city_lat))

        for index in range(len(city_info)):
            self.cities_dict[str(self.city_list[index])] = []
            self.strikes_dict[str(self.city_list[index])] = []
            self.reports_dict[str(self.city_list[index])] = []

        self.last_day = self.DB.read_db(begin, end)
        self.cities_dict, self.strikes_dict, self.total_posts, self.total_strikes, self.reports_dict = self.main_actions(self.last_day)
        self.filtered_reports_dict = self.dict_cleaning(self.reports_dict, uniq)
        self.cities_cleaned_dict, self.strikes_cleaned_dict = self.cities_and_strikes_cleaning(self.cities_dict, self.strikes_dict, uniq)
        self.is_build = True

    def build_map(self, is_update, begin, end, uniq):
        if is_update or not self.is_build:
            self.calculate(begin, end, uniq)
        self.map_creation(self.cities_cleaned_dict, self.strikes_cleaned_dict, self.total_posts, self.total_strikes, begin, end, uniq)

    def build_report(self, is_update, begin, end, uniq):
        if is_update or not self.is_build:
            self.calculate(begin, end, uniq)
        self.report(self.filtered_reports_dict, begin, end, uniq)
